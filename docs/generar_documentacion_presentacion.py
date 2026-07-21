from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "documentacion_tecnica_visual_presentacion.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 20, 20)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_font(run, size=None, bold=None, color=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", style=None, bold=False, italic=False, color=None, size=None, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color or INK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_font(r, color=INK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_font(r, color=INK)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, bold=True, color=DARK_BLUE)
        set_cell_shading(hdr[i], header_fill)
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(value))
            set_font(r, size=10, color=INK)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_callout(doc, title, text, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(text)
    set_font(r2, color=INK)
    set_table_width(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_figure(doc, rel_path, caption):
    path = ROOT / rel_path
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.9))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption)
    set_font(r, size=9, italic=True, color=MUTED)


def configure_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = sec.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Grupo 8 - Analisis de Datos ACP")
    set_font(r, size=9, color=MUTED)

    footer = sec.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Guia tecnica y visual para presentacion del codigo")
    set_font(r, size=9, color=MUTED)


def build_doc() -> None:
    doc = Document()
    configure_doc(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Guia tecnica y visual")
    set_font(r, size=26, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Proyecto: Analisis de Datos del Canal de Panama")
    set_font(r, size=15, color=MUTED)

    add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ["Curso / entrega", "Segundo Parcial - Pipeline + Visualizacion"],
            ["Objetivo del documento", "Explicar el codigo, el flujo tecnico, el dashboard y los resultados que se muestran."],
            ["Periodo analizado", "Octubre 2019 a diciembre 2025, con pronostico para enero a diciembre 2026."],
            ["Stack principal", "Python, pandas, NumPy, scikit-learn, matplotlib, seaborn, Streamlit y Plotly."],
        ],
        [1.55, 4.95],
        header_fill=LIGHT_GRAY,
    )

    add_callout(
        doc,
        "Idea central para exponer",
        "El proyecto toma datos de transitos del Canal de Panama, los limpia, los une con una segunda fuente macroeconomica, analiza tendencias, entrena un modelo predictivo y presenta todo en un dashboard interactivo.",
        fill=LIGHT_BLUE,
    )

    doc.add_page_break()

    doc.add_heading("1. Resumen del sistema", level=1)
    add_para(
        doc,
        "El repositorio esta organizado como un pipeline completo de analisis de datos. Cada carpeta representa una responsabilidad: ingesta, integracion, analisis exploratorio, modelado predictivo y visualizacion final en dashboard.",
    )
    add_table(
        doc,
        ["Modulo", "Responsabilidad", "Salida principal"],
        [
            ["persona1_ingesta", "Obtiene y limpia los transitos del Canal por fecha y segmento.", "canal_limpio.csv y canal_serie_mensual.csv"],
            ["persona2_pipeline", "Descarga/genera precio del petroleo y une ambas fuentes por mes.", "dataset_unificado.csv y dataset_unificado_completo.csv"],
            ["persona3_analisis", "Crea features, agregados, hallazgos y graficas exploratorias.", "canal_unificado.csv, insights.json y 10 figuras PNG"],
            ["persona4_modelo", "Entrena modelos de ML y pronostica transitos mensuales 2026.", "modelo_transitos.pkl, metricas y predicciones_2026.csv"],
            ["persona5_dashboard", "Muestra KPIs, tendencias, modelo, mapas y resumen con IA.", "Aplicacion Streamlit multipagina"],
        ],
        [1.45, 3.15, 1.9],
    )

    doc.add_heading("2. Flujo tecnico de datos", level=1)
    add_numbered(
        doc,
        [
            "Persona 1 genera la fuente base: tránsitos por segmento, fecha, calado, toneladas y peajes.",
            "Persona 2 toma la serie mensual total y le agrega el precio del barril desde el FMI o desde una muestra reproducible.",
            "Persona 3 usa el dataset por segmento para calcular eventos, estacionalidad, ranking, comparativas y figuras.",
            "Persona 4 usa la serie mensual total para crear features temporales, comparar modelos y producir el pronostico 2026.",
            "Persona 5 carga los CSV, JSON, PNG y el resumen del modelo para presentarlos en Streamlit.",
        ],
    )
    add_callout(
        doc,
        "Pregunta probable",
        "Si preguntan por que hay datos de muestra, la respuesta es: para que el pipeline sea reproducible y no dependa de internet o de una URL externa durante el desarrollo. Cuando se reemplaza por datos reales, el esquema de salida se mantiene.",
    )

    doc.add_heading("3. Guia para explicar el codigo", level=1)
    add_para(
        doc,
        "La forma mas clara de presentar el codigo es explicar cada archivo en cuatro partes: imports, constantes de ruta/configuracion, funciones de transformacion y bloque de ejecucion. Esa estructura se repite en casi todo el proyecto.",
    )
    add_table(
        doc,
        ["Parte del codigo", "Que significa", "Como explicarlo al profesor"],
        [
            ["Imports", "Librerias usadas: pandas, numpy, requests, scikit-learn, Streamlit, Plotly.", "Son herramientas externas que evitan programar desde cero lectura de datos, graficas, modelos y dashboard."],
            ["Constantes de rutas", "Variables como RUTA_BASE, RUTA_PROCESSED, OUTPUT, MODELS.", "Centralizan donde se leen y guardan archivos; si cambia una carpeta, se ajusta una variable."],
            ["Funciones", "Bloques reutilizables: cargar, limpiar, unir, entrenar, guardar.", "Dividen el problema en pasos pequeños, faciles de probar y explicar."],
            ["DataFrames", "Tablas de pandas llamadas df, df_canal, df_fuente2, serie, pred.", "Cada DataFrame representa una tabla temporal durante el pipeline."],
            ["main() / ejecutar()", "Orquestan el orden completo de cada modulo.", "Son el punto donde se conectan todas las funciones del archivo."],
            ["if __name__ == '__main__'", "Permite ejecutar el archivo como script.", "Si el archivo se importa desde otro modulo, no corre automaticamente todo el pipeline."],
        ],
        [1.45, 2.35, 2.7],
    )

    add_callout(
        doc,
        "Frase util",
        "El codigo esta escrito como pipeline funcional: cada funcion recibe datos, los transforma y devuelve un resultado. Despues el main une esas funciones en orden.",
        fill=LIGHT_BLUE,
    )

    doc.add_heading("3.1 Explicacion del flujo interno por archivo", level=2)
    add_table(
        doc,
        ["Archivo", "Flujo interno del codigo"],
        [
            ["ingesta_canal.py", "1) Decide modo de ingesta. 2) Carga o genera datos. 3) Limpia columnas, fechas, duplicados y negativos. 4) Agrupa por mes. 5) Guarda CSV."],
            ["ingesta_fuente2.py", "1) Intenta llamar la API del FMI. 2) Si falla, genera muestra. 3) Calcula variacion mensual y media movil. 4) Guarda copia local y copia para Persona 3."],
            ["pipeline.py", "1) Verifica que exista Fuente 1. 2) Ejecuta Fuente 2. 3) Hace merge mensual. 4) Hace merge por segmento. 5) Exporta datasets unificados."],
            ["preprocesamiento.py", "1) Carga canal_limpio. 2) Crea flags y ratios. 3) Une con combustible. 4) Calcula agregados para dashboard. 5) Guarda tablas procesadas."],
            ["analisis_tendencias.py", "1) Lee canal_unificado. 2) Calcula estadisticas, ranking, impacto, tendencia y estacionalidad. 3) Genera insights.json."],
            ["entrenamiento.py", "1) Prepara features. 2) Evalua modelos con TimeSeriesSplit. 3) Hace hold-out temporal. 4) Reentrena ganador. 5) Guarda modelo y metricas."],
            ["prediccion.py", "1) Carga modelo pickle. 2) Reconstruye features historicas. 3) Pronostica 12 meses recursivamente. 4) Guarda predicciones_2026.csv."],
            ["app.py y pages/*.py", "1) Configuran pagina Streamlit. 2) Cargan CSV/JSON con data_loader. 3) Construyen KPIs, tablas y graficas Plotly. 4) Renderizan pantalla."],
        ],
        [1.65, 4.85],
    )

    doc.add_heading("3.2 Operaciones de pandas que debes saber defender", level=2)
    add_table(
        doc,
        ["Operacion", "Donde aparece", "Explicacion sencilla"],
        [
            ["pd.read_csv(..., parse_dates=['fecha'])", "Carga de datasets", "Lee un CSV y convierte fecha a tipo datetime para poder ordenar, agrupar por mes y graficar en eje temporal."],
            ["df.copy()", "limpiar(), normalizar_canal(), unir_fuentes()", "Evita modificar accidentalmente el DataFrame original; se trabaja sobre una copia segura."],
            ["dropna(), fillna(), drop_duplicates()", "Limpieza de datos", "Elimina registros invalidos, rellena nulos numericos y quita filas repetidas."],
            ["groupby()", "Series mensuales, rankings y agregados", "Agrupa filas por fecha, segmento, periodo o anio para calcular sumas/promedios."],
            ["merge()", "Uniones entre fuentes", "Une dos tablas por una llave comun, como fecha o anio/mes."],
            ["rolling()", "Medias moviles y tendencia", "Calcula ventanas moviles, por ejemplo media de 3 meses o tendencia de 12 meses."],
            ["pct_change()", "Variacion del precio", "Calcula el cambio porcentual respecto al mes anterior."],
            ["np.select()", "Clasificacion de periodos", "Crea una columna categorica con reglas: sequia, recuperacion o baseline."],
        ],
        [1.75, 1.95, 2.8],
    )

    doc.add_heading("3.3 Codigo del modelo explicado paso a paso", level=2)
    add_table(
        doc,
        ["Bloque", "Que hace", "Por que importa"],
        [
            ["construir_features()", "Crea mes_sin, mes_cos, indice_tendencia, flags de regimen, lags y precio.", "Convierte una serie temporal en una tabla apta para scikit-learn."],
            ["lag_1", "Usa los transitos del mes anterior.", "Captura continuidad de corto plazo: un mes alto suele relacionarse con el siguiente."],
            ["lag_12", "Usa el mismo mes del anio anterior.", "Captura estacionalidad anual sin usar datos futuros."],
            ["TimeSeriesSplit", "Valida con ventanas expansivas respetando el tiempo.", "Evita entrenar con el futuro, que seria fuga de informacion."],
            ["mape()", "Calcula error porcentual absoluto medio.", "Es facil de explicar: error promedio en porcentaje."],
            ["pickle.dump()", "Guarda modelo, lista de features y nombre ganador.", "Permite usar el modelo despues sin volver a entrenar en cada pantalla."],
            ["pronostico recursivo", "Cada prediccion entra como historia para calcular el mes siguiente.", "Es necesario porque lag_1 y media_movil_3 dependen de meses previos."],
        ],
        [1.55, 2.55, 2.4],
    )

    doc.add_heading("3.4 Codigo del dashboard explicado paso a paso", level=2)
    add_table(
        doc,
        ["Bloque Streamlit / Plotly", "Que hace"],
        [
            ["st.set_page_config()", "Configura titulo, icono, ancho de pantalla y comportamiento inicial del dashboard."],
            ["utils/data_loader.py", "Centraliza todas las lecturas de CSV, JSON y rutas de imagen para que las paginas no repitan rutas."],
            ["try / except en app.py", "Si falta un archivo o falla la carga, muestra un error claro y detiene la app con st.stop()."],
            ["st.columns()", "Divide la pantalla en columnas para KPIs, graficas o tablas lado a lado."],
            ["st.metric()", "Muestra indicadores resumidos como promedio de transitos, peajes, calado y modelo ganador."],
            ["st.tabs()", "Organiza muchas visualizaciones sin saturar una sola pantalla."],
            ["go.Figure() + add_trace()", "Construye graficas con capas: historico, pronostico, reales y predichos."],
            ["px.bar(), px.pie(), px.imshow()", "Genera graficos rapidos de barras, pastel y mapas de calor con Plotly Express."],
            ["st.session_state", "Guarda el resumen generado por LLM para mantenerlo visible despues de presionar el boton."],
            ["st.download_button()", "Permite descargar el resumen ejecutivo en Markdown."],
        ],
        [2.25, 4.25],
    )

    doc.add_heading("3.5 Preguntas de codigo y respuestas cortas", level=2)
    add_table(
        doc,
        ["Si preguntan...", "Respuesta tecnica breve"],
        [
            ["¿Por que usan Path y rutas absolutas?", "Para resolver archivos desde cualquier lugar donde se ejecute el script y evitar errores por directorio actual."],
            ["¿Por que hay funciones separadas?", "Porque cada funcion tiene una responsabilidad: cargar, limpiar, transformar, entrenar o guardar."],
            ["¿Que es un DataFrame?", "Una tabla en memoria con columnas y filas, similar a Excel, pero manipulable con codigo."],
            ["¿Que significa persistir?", "Guardar el resultado en disco, normalmente como CSV, JSON, PNG, PKL o DOCX."],
            ["¿Por que se usa left join?", "Para conservar la fuente principal aunque falten datos de la segunda fuente."],
            ["¿Por que interpolan precios faltantes?", "Porque si un mes no trae precio, se estima entre valores vecinos para no romper el modelo ni el dashboard."],
            ["¿Por que se descartan 12 filas del modelo?", "Porque lag_12 necesita un anio previo; los primeros 12 meses no tienen ese dato."],
            ["¿Que es feature_importances_?", "Es la importancia relativa que modelos de arbol asignan a cada variable predictora."],
        ],
        [2.3, 4.2],
    )

    doc.add_page_break()
    doc.add_heading("4. Persona 1: ingesta ACP", level=1)
    add_para(doc, "Archivo principal: persona1_ingesta/src/ingesta_canal.py.")
    add_bullets(
        doc,
        [
            "Modos de entrada: url, local y muestra.",
            "Normaliza nombres de columnas, fechas, duplicados, nulos numericos y transitos negativos.",
            "Construye una serie mensual total agregando todos los segmentos.",
            "Entrega 750 filas por segmento y mes, mas una serie mensual de 75 meses.",
        ],
    )
    add_table(
        doc,
        ["Funcion", "Que hace"],
        [
            ["ingestar()", "Decide si los datos vienen de URL, archivo local o generador de muestra."],
            ["limpiar()", "Estandariza columnas, convierte fecha, elimina duplicados y valida negativos."],
            ["construir_serie_mensual()", "Agrupa por mes y suma tránsitos para crear la serie total."],
            ["guardar()", "Persiste los CSV en data/processed."],
        ],
        [2.1, 4.4],
    )

    doc.add_heading("5. Persona 2: integracion de fuentes", level=1)
    add_para(doc, "Archivos principales: persona2_pipeline/src/ingesta_fuente2.py y persona2_pipeline/src/pipeline.py.")
    add_bullets(
        doc,
        [
            "La segunda fuente es el precio mensual del petroleo crudo del FMI, indicador POILAPSP.",
            "Se calculan var_mensual_pct y precio_barril_usd_ma3 para enriquecer el analisis y el modelo.",
            "El join se hace por fecha o por anio/mes, segun el nivel de agregacion.",
            "Si hay valores faltantes despues del join, se interpolan o se rellenan de forma controlada.",
        ],
    )
    add_table(
        doc,
        ["Dataset", "Forma", "Uso"],
        [
            ["dataset_unificado.csv", "75 filas x 7 columnas", "Entrada principal del modelo predictivo."],
            ["dataset_unificado_completo.csv", "750 filas x 11 columnas", "Analisis por segmento en Persona 3."],
            ["fuente2_combustibles.csv", "fecha, anio, mes, precio_barril_usd", "Copia compatible con el preprocesamiento de Persona 3."],
        ],
        [2.1, 1.6, 2.8],
    )

    doc.add_heading("6. Persona 3: analisis exploratorio", level=1)
    add_para(doc, "Archivos principales: preprocesamiento.py, analisis_tendencias.py y visualizaciones.py.")
    add_bullets(
        doc,
        [
            "Normaliza la fuente del canal y crea fase_fiscal, periodo_sequia y periodo_recuperacion.",
            "Calcula ratio_toneladas_por_transito y peaje_por_tonelada_usd para enriquecer el analisis.",
            "Genera estadisticas descriptivas, ranking de segmentos, impacto de la sequia, tendencia anual y descomposicion estacional.",
            "Produce 10 figuras PNG que luego se pueden mostrar en el dashboard.",
        ],
    )
    add_table(
        doc,
        ["Hallazgo", "Valor para decir en presentacion"],
        [
            ["Cobertura", "75 meses, 10 segmentos, 750 observaciones."],
            ["Total de transitos", "72,543 transitos entre 2019-10 y 2025-12."],
            ["Segmento lider", "Portacontenedores con 24.45% del total."],
            ["Sequía", "Pasajeros fue el segmento mas golpeado: -28.49% contra baseline."],
            ["Recuperacion 2025", "Quimiqueros lidera el rebote: +22.18% contra baseline."],
            ["Estacionalidad", "Pico en abril (+8.66%) y valle en agosto (-9.26%)."],
        ],
        [1.8, 4.7],
    )

    doc.add_heading("7. Persona 4: modelo predictivo", level=1)
    add_para(doc, "Archivos principales: preparacion_datos.py, entrenamiento.py, prediccion.py y visualizaciones.py.")
    add_table(
        doc,
        ["Tipo de feature", "Variables", "Motivo"],
        [
            ["Calendario", "mes_sin, mes_cos, indice_tendencia", "Capturan estacionalidad y tendencia temporal."],
            ["Regimen", "periodo_sequia, periodo_recuperacion", "Representan quiebres operativos relevantes."],
            ["Autorregresivas", "lag_1, lag_12, media_movil_3", "Usan memoria de corto y largo plazo de la serie."],
            ["Exogenas", "precio_barril_usd, precio_barril_usd_ma3", "Agregan contexto macroeconomico externo."],
        ],
        [1.5, 2.4, 2.6],
    )
    add_para(
        doc,
        "El entrenamiento compara Regresion Lineal, Random Forest, Gradient Boosting y un baseline Naive Estacional. La seleccion se hace por menor MAPE en validacion cruzada temporal con TimeSeriesSplit de 5 folds.",
    )
    add_table(
        doc,
        ["Modelo", "MAPE CV", "MAPE hold-out", "Lectura"],
        [
            ["Gradient Boosting", "6.94%", "14.79%", "Modelo ganador por CV temporal."],
            ["Random Forest", "11.42%", "14.04%", "Buen competidor no lineal."],
            ["Regresion Lineal", "26.26%", "13.26%", "Interpretable, pero peor en CV."],
            ["Naive Estacional", "N/A", "23.66%", "Baseline simple t = t-12."],
        ],
        [1.8, 1.1, 1.2, 2.4],
    )
    add_callout(
        doc,
        "Como responder sobre el R2 negativo",
        "El hold-out 2024-10 a 2025-12 funciona como prueba de estres porque contiene el salto de nivel de 2025. El R2 negativo indica que ese quiebre fue dificil de extrapolar, pero la seleccion del modelo se basa en CV temporal, que es mas estable para una serie corta.",
    )
    add_table(
        doc,
        ["Pronostico 2026", "Valor"],
        [
            ["Total anual proyectado", "13,565 transitos."],
            ["Meses mas altos", "Enero a abril, cerca de 1,174 a 1,179 transitos."],
            ["Meses mas bajos", "Agosto a octubre, cerca de 1,071 a 1,076 transitos."],
            ["Supuesto de precio", "Precio del barril plano con el ultimo valor observado."],
        ],
        [2.0, 4.5],
    )

    doc.add_page_break()
    doc.add_heading("8. Persona 5: dashboard y experiencia visual", level=1)
    add_para(doc, "El dashboard esta hecho con Streamlit, Plotly y una capa de carga centralizada en persona5_dashboard/src/utils/data_loader.py.")
    add_table(
        doc,
        ["Pantalla", "Que muestra", "Que explicar"],
        [
            ["Inicio", "KPIs, serie historica + pronostico, insights y comparativa de periodos.", "Es la vista ejecutiva para entender el proyecto rapido."],
            ["Tendencias", "Serie temporal, ranking, pie chart, estacionalidad, heatmaps y figuras EDA.", "Sirve para demostrar patrones historicos y efecto de la sequia."],
            ["Modelo Predictivo", "Modelo ganador, MAPE CV, pronostico 2026, metricas e importancia de features.", "Conecta el analisis con prediccion y validacion."],
            ["Resumen LLM", "Genera reportes con OpenAI, Anthropic o modo local.", "Convierte datos y hallazgos en texto ejecutivo."],
            ["Mapas", "Ubicacion del canal, ruta y visualizaciones por zona/segmento.", "Apoya la explicacion geografica del fenomeno."],
        ],
        [1.4, 3.0, 2.1],
    )
    add_bullets(
        doc,
        [
            "Plotly se usa para interactividad: hover, zoom, filtros visuales y graficos responsivos.",
            "Las metricas de Streamlit resumen variables clave sin obligar al usuario a leer tablas largas.",
            "Las figuras PNG de Persona 3 y 4 funcionan como evidencia visual reproducible del pipeline.",
            "El CSS customizado en app.py estiliza tarjetas de metricas con fondo oscuro y texto claro.",
        ],
    )

    doc.add_heading("9. Figuras clave para defender", level=1)
    add_figure(doc, "persona3_analisis/figures/01_serie_mensual.png", "Figura 1. Serie mensual: muestra la caida por sequia y la recuperacion posterior.")
    add_figure(doc, "persona3_analisis/figures/06_comparativa_periodos.png", "Figura 2. Comparativa baseline, sequia y recuperacion por segmento.")
    add_figure(doc, "persona4_modelo/figures/03_importancia_features.png", "Figura 3. Importancia de features: periodo_sequia domina la prediccion.")
    add_figure(doc, "persona4_modelo/figures/04_pronostico_2026.png", "Figura 4. Pronostico 2026: historico y proyeccion mensual.")

    doc.add_heading("10. Preguntas probables del profesor", level=1)
    add_table(
        doc,
        ["Pregunta", "Respuesta sugerida"],
        [
            ["¿Cual es la variable objetivo del modelo?", "transitos_totales mensuales del Canal de Panama."],
            ["¿Por que no se separo train/test aleatoriamente?", "Porque es una serie temporal; mezclar fechas permitiria entrenar con informacion del futuro."],
            ["¿Que significa MAPE?", "Error porcentual absoluto medio. Un MAPE CV de 6.94% indica que, en promedio, el error relativo fue cercano a 6.94% en validacion temporal."],
            ["¿Por que usar Gradient Boosting?", "Porque fue el mejor por MAPE en validacion cruzada temporal y captura relaciones no lineales entre regimenes, lags y calendario."],
            ["¿Que hace lag_12?", "Usa el valor del mismo mes del año anterior, capturando estacionalidad anual."],
            ["¿Que pasa si no hay internet?", "El pipeline tiene modos de muestra y fallback para no bloquear la ejecucion ni la demo."],
            ["¿La correlacion con peajes prueba causalidad?", "No. Peajes y transitos estan casi perfectamente correlacionados porque peajes se generan como funcion creciente de transitos en la muestra."],
            ["¿Que limitacion principal tiene el modelo?", "La serie modelable es corta, 63 meses, y el salto de 2025 es un quiebre dificil de extrapolar."],
            ["¿Como se actualiza el dashboard?", "Regenerando los CSV/JSON/PNG con los pipelines; Streamlit los vuelve a cargar desde data_loader.py."],
            ["¿Que muestra la pagina de LLM?", "Toma contexto de estadisticas, insights, modelo y predicciones para generar resumen ejecutivo con API o modo local."],
        ],
        [2.1, 4.4],
    )

    doc.add_heading("11. Como ejecutar y presentar", level=1)
    add_table(
        doc,
        ["Paso", "Comando / accion"],
        [
            ["1", "python persona1_ingesta/src/ingesta_canal.py --modo muestra"],
            ["2", "python persona2_pipeline/src/pipeline.py"],
            ["3", "python persona3_analisis/src/run_pipeline.py"],
            ["4", "python persona4_modelo/src/run_pipeline.py"],
            ["5", "streamlit run persona5_dashboard/src/app.py"],
        ],
        [0.7, 5.8],
    )
    add_callout(
        doc,
        "Orden recomendado de exposicion",
        "Primero explica el objetivo y el flujo completo. Luego muestra una carpeta por persona. Despues abre el dashboard y conecta cada grafico con el archivo CSV/JSON/PNG que lo alimenta. Cierra con metricas del modelo y limitaciones.",
        fill=LIGHT_BLUE,
    )

    doc.add_heading("12. Limitaciones y mejoras", level=1)
    add_bullets(
        doc,
        [
            "Los datos de muestra son utiles para desarrollo, pero los resultados finales deben recalcularse con fuentes reales confirmadas.",
            "El pronostico 2026 usa precio del barril plano; se puede mejorar integrando un pronostico externo de combustible.",
            "El modelo predice tránsitos totales; una mejora seria entrenar modelos por segmento.",
            "El dashboard podria enlazar explicitamente la pagina de mapas en la navegacion principal si se desea mostrarla siempre.",
            "Se pueden agregar pruebas automatizadas para validar esquemas de CSV y evitar errores de columnas faltantes.",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
